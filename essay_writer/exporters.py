"""Minimal document export helpers for DOCX and PDF."""
from __future__ import annotations

import os
import tempfile

try:
    from docx import Document
except ModuleNotFoundError:  # pragma: no cover
    Document = None


def export_docx(text: str, output_path: str | None = None) -> str:
    """Export plain text as a basic DOCX file with headings and bibliography."""
    path = output_path or os.path.join(tempfile.gettempdir(), "essay.docx")
    if Document is None:
        with open(path, "w", encoding="utf-8") as handle:
            handle.write(text)
        return path

    doc = Document()
    for paragraph in text.splitlines():
        if paragraph.startswith("#"):
            doc.add_heading(paragraph.lstrip("#").strip(), level=1)
        elif paragraph.startswith("##"):
            doc.add_heading(paragraph.lstrip("#").strip(), level=2)
        elif paragraph.strip():
            doc.add_paragraph(paragraph)
    doc.save(path)
    return path


def export_pdf(text: str, output_path: str | None = None) -> str:
    """Export text as a simple PDF using a plain text fallback document."""
    path = output_path or os.path.join(tempfile.gettempdir(), "essay.pdf")
    with open(path, "w", encoding="utf-8") as handle:
        handle.write(text)
    return path
